import json
import os
import base64
from typing import Dict, Any
import PyPDF2
import docx
from io import BytesIO

def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    '''
    Business: Обрабатывает загрузку и извлечение текста из PDF, Word и TXT файлов
    Args: event - dict с httpMethod, body (file_data, file_name, file_type)
          context - объект с request_id, function_name и другими атрибутами  
    Returns: HTTP ответ с извлеченным текстом
    '''
    method: str = event.get('httpMethod', 'GET')
    
    # Обработка CORS OPTIONS запроса
    if method == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'POST, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type, Authorization',
                'Access-Control-Max-Age': '86400'
            },
            'body': ''
        }
    
    if method != 'POST':
        return {
            'statusCode': 405,
            'headers': {'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': 'Method not allowed'})
        }
    
    try:
        # Парсим данные запроса
        body_data = json.loads(event.get('body', '{}'))
        file_data_b64 = body_data.get('file_data', '')
        file_name = body_data.get('file_name', '')
        file_type = body_data.get('file_type', '')
        
        if not file_data_b64 or not file_name:
            return {
                'statusCode': 400,
                'headers': {'Access-Control-Allow-Origin': '*'},
                'body': json.dumps({'error': 'File data and name are required'})
            }
        
        # Декодируем base64 данные
        try:
            file_data = base64.b64decode(file_data_b64)
        except Exception as e:
            return {
                'statusCode': 400,
                'headers': {'Access-Control-Allow-Origin': '*'},
                'body': json.dumps({'error': 'Invalid base64 file data'})
            }
        
        # Проверяем размер файла (максимум 50MB)
        if len(file_data) > 50 * 1024 * 1024:
            return {
                'statusCode': 400,
                'headers': {'Access-Control-Allow-Origin': '*'},
                'body': json.dumps({'error': 'File too large. Maximum size: 50MB'})
            }
        
        # Извлекаем текст в зависимости от типа файла
        extracted_text = ""
        
        if file_name.lower().endswith('.pdf'):
            extracted_text = extract_text_from_pdf(file_data)
        elif file_name.lower().endswith(('.doc', '.docx')):
            extracted_text = extract_text_from_word(file_data)
        elif file_name.lower().endswith('.txt'):
            extracted_text = extract_text_from_txt(file_data)
        else:
            return {
                'statusCode': 400,
                'headers': {'Access-Control-Allow-Origin': '*'},
                'body': json.dumps({'error': 'Unsupported file type. Supported: PDF, DOC, DOCX, TXT'})
            }
        
        if not extracted_text.strip():
            return {
                'statusCode': 400,
                'headers': {'Access-Control-Allow-Origin': '*'},
                'body': json.dumps({'error': 'No text could be extracted from the file'})
            }
        
        # Формируем ответ
        result = {
            'text': extracted_text,
            'file_name': file_name,
            'file_size': len(file_data),
            'text_length': len(extracted_text),
            'request_id': context.request_id
        }
        
        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*'
            },
            'body': json.dumps(result, ensure_ascii=False)
        }
        
    except json.JSONDecodeError:
        return {
            'statusCode': 400,
            'headers': {'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': 'Invalid JSON in request body'})
        }
    except Exception as e:
        print(f"Error processing document: {str(e)}")
        return {
            'statusCode': 500,
            'headers': {'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': 'Internal server error'})
        }


def extract_text_from_pdf(file_data: bytes) -> str:
    """Извлекает текст из PDF файла"""
    try:
        file_stream = BytesIO(file_data)
        pdf_reader = PyPDF2.PdfReader(file_stream)
        
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")


def extract_text_from_word(file_data: bytes) -> str:
    """Извлекает текст из Word файла (.doc, .docx)"""
    try:
        file_stream = BytesIO(file_data)
        doc = docx.Document(file_stream)
        
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Также извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting text from Word document: {str(e)}")


def extract_text_from_txt(file_data: bytes) -> str:
    """Извлекает текст из TXT файла"""
    try:
        # Пробуем разные кодировки
        encodings = ['utf-8', 'cp1251', 'latin1', 'ascii']
        
        for encoding in encodings:
            try:
                text = file_data.decode(encoding)
                return text.strip()
            except UnicodeDecodeError:
                continue
        
        # Если ни одна кодировка не сработала, используем errors='ignore'
        text = file_data.decode('utf-8', errors='ignore')
        return text.strip()
        
    except Exception as e:
        raise Exception(f"Error extracting text from TXT file: {str(e)}")